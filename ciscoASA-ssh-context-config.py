# -*- coding: utf-8 -*-
"""
Created on Wed Nov 30 17:10:30 2016

@author: root
"""


import paramiko
import re
import time
import datetime
import docx 
from termcolor import colored

context = []
tcp_ssh = 22
log_path = "/home/user/Документы/log.docx"
pathConf = "/home/user/Документы/config/"
now = datetime.datetime.now()
date = now.strftime("%d.%m.%Y %I:%M %p")  # получаем текущее время
print(now.strftime("%d.%m.%Y %I:%M:%S %p"))  # выводим текущее время


def log_conf_save_succes(log_path, hostname):
    doc = docx.Document(log_path)
    now = datetime.datetime.now()
    date = now.strftime("%d.%m.%Y")
    doc.add_paragraph("= Device configuration " + hostname + " successfully saved " + "- " + date, style='List')
    doc.save(log_path)
    
    
def enable_pass_incorrect(log_path, hostname):
    doc = docx.Document(log_path)
    doc.add_heading(u"= ОШИБКА!!! Пароль enable неверен на " + hostname + "=", level=2)
    doc.save(log_path)


def login_pass_incorrect(log_path, hostname):
    doc = docx.Document(log_path)
    doc.add_heading(u"=ОШИБКА!!! Неверный Логин/пароль на " + hostname + "=", level=2)
    doc.save(log_path)


def asa_login_ssh(hostname='192.168.252.150', user='cisco', password='cisco', enablepass='cisco', tcp_ssh=22):
    ssh = paramiko.SSHClient() 
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())  # политика принятия ssh  ключа при подключении к устройству.
    
    try:
        a = ssh.connect(hostname, username=user, password=password, port=tcp_ssh, timeout=10)  #  подключение к утройству
        chan = ssh.invoke_shell()  # работа  с  со строкой ssh в случае испрользования нескольких команд
        time.sleep(1)  # приостановить выполнение программы на заданное количество секунд
        b = chan.recv(100).split("\r")  #  в переменную b помещаем вывод консоли, после установления соединения
        c = re.search(">", str(b)) # ищет в строке символ ">" функция возвращает, то где символ находится, если он есть и None, если его нет

        if c is not None:  # если данный символ ">"  есть выполняем действия ниже
           chan.send('enable' + '\n')  # вводим в командную строку enable
           chan.send(enablepass + '\n')  # вводим в командную строку пароль на enable
           time.sleep(1)  # приостанавливает выполнение программы на заданное количество секунд. неоходимо для записи содержимого консоли в переменные.
           b = chan.recv(100)

           if (re.search("#", b)) is not None:
              chan.send("show mode" + "\n")
              time.sleep(3)
              c = chan.recv(100).split(' ')
			  
              """ determiner  the mode operation (multiple or single)"""
              if (re.search('multiple', str(c))) is not None:
                 """action in multiple mode"""
                 print(colored("ASA Multiple mode \n", "red"))
                 chan.send("changeto conte sys"+"\n")  # changeto system context
                 chan.send("show run | in hostnam" + "\n")
                 time.sleep(1)
                 c = chan.recv(10000)
                 device = (str(re.search('hostname \w+', c).group(0)).split(' '))[1]  # ищем в строке "с" с помощью функции re.search регулярное выражение "hostname \w+", где должно содержаться слово  hostname + еще слово за ним. с помощью .group(0)  выдераем значение которое попало под сопадение(так как без этой функция re.search отдает адрес в памяти где было совпадение), далее с помощью функции  split  делин полученную строку на отдельные слова и помещаем в список из этого списка берем второй элемент
                 print(device)
                 chan.send("show run | in conte" + "\n")  # view contexts list
                 time.sleep(3)
                 c = chan.recv(10000).split('\n')
                 for i in range(len(c)):
                     test = re.findall(r'^context \w+', c[i])
                     if len(test) != 0:
                         print(test)
                         context.append(str(test).strip('[').strip(']').strip('\''))  # populate list with name contexts             
                 print(colored("Количество контекстов= " + str(len(context)), "yellow"))
                 print(context)
				 
                 DevConfFile = str(pathConf) + str(device) + "-" + str("context system") + '-' + date + str(".txt")  # form the full path where context configuration will be store

                 print(DevConfFile)

                 # save configuration system context
                 chan.send(" terminal pager 0" + "\n")
                 chan.send("show run" + "\n")
                 time.sleep(11)
                 b = han.recv(20000) 
                 config = open(DevConfFile, "w")
                 config.write(b)
                 config.close()
                 log_conf_save_succes(log_path, hostname)
                 print(colored("Конфигурация успешно сохранена \n", "green"))
				 
                 for i in range(len(context)):
                     """save configuration of all contexts"""
                     chan.send("changeto " + context[i] + "\n")
                     chan.send("show run | in hostnam" + "\n")
                     time.sleep(1)
                     c=chan.recv(10000)
                     DevConfFile = str(pathConf) + str(context[i]) + " - " + date + str(".txt")
					 
                     print(DevConfFile)
					 
                     chan.send(" terminal pager 0" + "\n")
                     chan.send("show run" + "\n")
                     time.sleep(11)
                     b = chan.recv(20000)
                     config = open(DevConfFile, "w")
                     config.write(b)
                     config.close()
                     log_conf_save_succes(log_path, hostname)
                     print(colored("Конфигурация " + context[i] + " успешно сохранена \n", "green"))
              
              elif (re.search('single', str(c))) is not None:
                  """ action in case single mode"""
                  print(colored("ASA Single mode \n", "red"))
                  chan.send("show run | in hostnam" + "\n")
                  time.sleep(1)
                  c=chan.recv(10000)
                  device=(str(re.search('hostname \w+', c).group(0)).split(' '))[1]# ищем в строке "с" с помощью функции re.search регулярное выражение "hostname \w+", где должно содержаться слово  hostname + еще слово за ним. с помощью .group(0)  выдераем значение которое попало под сопадение(так как без этой функция re.search отдает адрес в памяти где было совпадение), далее с помощью функции  split  делин полученную строку на отдельные слова и помещаем в список из этого списка берем второй элемент
                  print(device)
                  DevConfFile=str(pathConf) + str(device) + ' - ' + date + str(".txt")
                  print(DevConfFile)
                  chan.send(" terminal pager 0" + "\n")
                  chan.send("show run" + "\n")
                  time.sleep(11)
                  b = chan.recv(20000)
                  config = open(DevConfFile, "w")
                  config.write(b)
                  config.close()
                  log_conf_save_succes(log_path, hostname)
                  print(colored("Конфигурация " + hostname + " успешно сохранена \n", "green"))
              
           else:
             enable_pass_incorrect(log_path, hostname)
             print (colored("=Пароль enable неверен= \n", "red"))
        else:
             DevConfFile = str(pathConf) + str(device) + "-" + date + str(".txt")
             chan.send("terminal len 0" + "\n")
             chan.send("show run" + "\n")
             time.sleep(11)
             b = chan.recv(20000)
             config = open(DevConfFile, "w")
             config.write(b)
             config.close()
    except paramiko.AuthenticationException:
        login_pass_incorrect(log_path, hostname)
        print (colored ( "=Неверный Логин/пароль на  " + hostname, "red"))
             

asa_login_ssh(hostname='192.168.252.150', user='cisco', password='cisco', enablepass='cisc', tcp_ssh=22)